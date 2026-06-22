# -*- coding: utf-8 -*-
"""Generate the 27 Markets COMPLETE WORKFLOW specification as a .docx (24 sections)."""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image

BASE = r'c:/Users/gowth/OneDrive/Desktop/27 trading'
DIAG = os.path.join(BASE, 'assets', 'diagrams')
LOGO = r'C:/Users/gowth/Downloads/corecare consultancy.png'
WM = os.path.join(BASE, 'assets', 'watermark.png')

def make_watermark(src, dst, opacity=0.10):
    """Create a faint, washed-out version of the logo for use as a page watermark."""
    im = Image.open(src).convert('RGBA')
    bg = Image.new('RGBA', im.size, (255, 255, 255, 255))
    comp = Image.alpha_composite(bg, im)          # flatten logo onto white
    faded = Image.blend(bg, comp, opacity)        # mostly-white => faint logo
    faded.convert('RGB').save(dst)

def set_watermark(section, image_path, width_in=4.6):
    """Anchor a centered, behind-text image in the section header so it repeats on every page."""
    header = section.header
    header.is_linked_to_previous = False
    run = header.add_paragraph().add_run()
    run.add_picture(image_path, width=Inches(width_in))
    drawing = run._r.find(qn('w:drawing'))
    inline = drawing.find(qn('wp:inline'))
    extent = inline.find(qn('wp:extent'))
    cx, cy = extent.get('cx'), extent.get('cy')
    docPr = inline.find(qn('wp:docPr'))
    graphic = inline.find(qn('a:graphic'))
    anchor = OxmlElement('wp:anchor')
    for k, v in {'distT':'0','distB':'0','distL':'0','distR':'0','simplePos':'0',
                 'relativeHeight':'251650048','behindDoc':'1','locked':'0',
                 'layoutInCell':'1','allowOverlap':'1'}.items():
        anchor.set(k, v)
    sp = OxmlElement('wp:simplePos'); sp.set('x','0'); sp.set('y','0'); anchor.append(sp)
    ph = OxmlElement('wp:positionH'); ph.set('relativeFrom','page')
    a1 = OxmlElement('wp:align'); a1.text='center'; ph.append(a1); anchor.append(ph)
    pv = OxmlElement('wp:positionV'); pv.set('relativeFrom','page')
    a2 = OxmlElement('wp:align'); a2.text='center'; pv.append(a2); anchor.append(pv)
    ex = OxmlElement('wp:extent'); ex.set('cx',cx); ex.set('cy',cy); anchor.append(ex)
    ee = OxmlElement('wp:effectExtent')
    for k in ('l','t','r','b'): ee.set(k,'0')
    anchor.append(ee)
    anchor.append(OxmlElement('wp:wrapNone'))
    anchor.append(docPr)
    anchor.append(OxmlElement('wp:cNvGraphicFramePr'))
    anchor.append(graphic)
    drawing.remove(inline); drawing.append(anchor)

RED = RGBColor(0xE1, 0x1D, 0x2E)
DARK = RGBColor(0x11, 0x11, 0x11)
GRAY = RGBColor(0x55, 0x55, 0x55)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

doc = Document()
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(6)

def shade(cell, hexc):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),hexc)
    tcPr.append(shd)

def h(text, level=1):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.color.rgb = RED if level == 1 else DARK
    return p

def para(text, bold=False, italic=False, color=None, size=None, after=6):
    p = doc.add_paragraph(); r = p.add_run(text)
    r.bold=bold; r.italic=italic
    if color: r.font.color.rgb=color
    if size: r.font.size=Pt(size)
    p.paragraph_format.space_after=Pt(after); return p

def bullet(text, prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if prefix:
        r=p.add_run(prefix); r.bold=True; p.add_run(text)
    else: p.add_run(text)
    return p

def step(text):
    return doc.add_paragraph(text, style='List Number')

def flow(text):
    p = doc.add_paragraph(); r = p.add_run(text)
    r.font.name='Consolas'; r.font.size=Pt(9); r.font.color.rgb=GRAY
    return p

def add_field(paragraph, field_code):
    """Insert a Word field (e.g. PAGE, TOC) that updates on open."""
    r1 = paragraph.add_run(); fc1 = OxmlElement('w:fldChar'); fc1.set(qn('w:fldCharType'),'begin'); r1._r.append(fc1)
    r2 = paragraph.add_run(); it = OxmlElement('w:instrText'); it.set(qn('xml:space'),'preserve'); it.text=field_code; r2._r.append(it)
    r3 = paragraph.add_run(); fc2 = OxmlElement('w:fldChar'); fc2.set(qn('w:fldCharType'),'separate'); r3._r.append(fc2)
    r4 = paragraph.add_run('?')
    r5 = paragraph.add_run(); fc3 = OxmlElement('w:fldChar'); fc3.set(qn('w:fldCharType'),'end'); r5._r.append(fc3)
    return r4

def add_page_footer():
    sec = doc.sections[0]
    fp = sec.footer.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.add_run('27 Markets — Complete Workflow   |   Core Care Consultancy   |   Page ').font.size = Pt(8)
    add_field(fp, 'PAGE').font.size = Pt(8)
    fp.add_run(' of ').font.size = Pt(8)
    add_field(fp, 'NUMPAGES').font.size = Pt(8)

def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    t.style='Light Grid Accent 1'
    for i,head in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=''
        run=c.paragraphs[0].add_run(head); run.bold=True; run.font.color.rgb=WHITE; run.font.size=Pt(9)
        shade(c,'E11D2E')
    for row in rows:
        cells=t.add_row().cells
        for i,v in enumerate(row):
            cells[i].text=''; run=cells[i].paragraphs[0].add_run(str(v)); run.font.size=Pt(9)
    if widths:
        for i,w in enumerate(widths):
            for r_ in t.rows: r_.cells[i].width=Inches(w)
    doc.add_paragraph(); return t

def img(path, width=6.3):
    if not os.path.exists(path): return None
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Inches(width)); return p

def caption(text):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(text); r.italic=True; r.font.size=Pt(8.5); r.font.color.rgb=GRAY; return p

# ================= COVER =================
if os.path.exists(LOGO):
    lp=doc.add_paragraph(); lp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    lp.add_run().add_picture(LOGO, width=Inches(1.7))
t=doc.add_paragraph(); t.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=t.add_run('27 MARKETS'); r.bold=True; r.font.size=Pt(34); r.font.color.rgb=RED
s=doc.add_paragraph(); s.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=s.add_run('Trade Beyond Limits'); r.italic=True; r.font.size=Pt(13); r.font.color.rgb=GRAY
s2=doc.add_paragraph(); s2.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=s2.add_run('Complete Platform Workflow Specification'); r.bold=True; r.font.size=Pt(14)
m=doc.add_paragraph(); m.alignment=WD_ALIGN_PARAGRAPH.CENTER
m.add_run('Version 1.0  |  2026-06-22  |  CFD/Forex Brokerage — Marketing Site + Client Portal').font.size=Pt(10)
doc.add_paragraph()
i=doc.add_paragraph(); i.alignment=WD_ALIGN_PARAGRAPH.CENTER
ir=i.add_run('End-to-end operational workflow covering actors, modules, user journeys, money movement, '
             'compliance, statuses, data model, permissions, and integrations.')
ir.italic=True; ir.font.color.rgb=GRAY; ir.font.size=Pt(10.5)
doc.add_paragraph()
pb=doc.add_paragraph(); pb.alignment=WD_ALIGN_PARAGRAPH.CENTER
pr=pb.add_run('Prepared by Core Care Consultancy'); pr.bold=True; pr.font.size=Pt(11); pr.font.color.rgb=DARK
doc.add_page_break()

# ================= TABLE OF CONTENTS =================
add_page_footer()
toc_h = doc.add_paragraph(); tr = toc_h.add_run('Table of Contents'); tr.bold=True; tr.font.size=Pt(16); tr.font.color.rgb=RED
toc_p = doc.add_paragraph()
add_field(toc_p, 'TOC \\o "1-2" \\h \\z \\u')
doc.add_page_break()

# ===== 1. ACTORS =====
h('1. Application Actors / User Roles', 1)
table(['Actor','Scope','Description'], [
    ['Visitor (Lead)','Public','Anonymous prospect browsing the marketing site; can submit lead/demo/contact forms.'],
    ['Registered Client','Portal','Verified or pending account holder; trades, funds, manages account.'],
    ['IB / Partner','Partner portal','Introducing Broker / affiliate earning rebates on referred client volume.'],
    ['Support Agent','Admin','Handles tickets, live chat, basic client assistance.'],
    ['KYC / Compliance Officer','Admin','Reviews and approves/rejects KYC, AML checks, risk flags.'],
    ['Finance Officer','Admin','Approves deposits/withdrawals/transfers, reconciles funds.'],
    ['Administrator','Admin','Full configuration: accounts, leverage, content, roles, partners.'],
    ['System / Scheduler','Automated','Notifications, status transitions, rebate calc, scheduled jobs.'],
], widths=[1.6,1.3,3.5])

# ===== 2. MODULES =====
h('2. Full System Modules', 1)
bullet('Home, About, Trading Accounts, Markets, Partnership/IB, Contact, Login, Register.', 'Public Website: ')
bullet('Multi-step registration, demo request, lead capture, email verification.', 'Onboarding: ')
bullet('Identity, Proof of Address, Selfie — upload, review, status.', 'KYC / Compliance: ')
bullet('Dashboard, Accounts, account creation (Standard/Raw/VIP), leverage selection.', 'Account Management: ')
bullet('Deposit, Withdraw, Internal Transfer, Transaction History.', 'Funds / Cashier: ')
bullet('Platform installers + document download center.', 'Downloads: ')
bullet('Tickets, live chat, FAQ.', 'Support / Helpdesk: ')
bullet('IB registration, referral links, rebate reports, marketing assets.', 'Partner / IB: ')
bullet('Profile, security, preferences, notification settings.', 'Profile: ')
bullet('User mgmt, KYC queue, finance approvals, content/CMS, partner mgmt, audit log.', 'Admin Back-office: ')
bullet('In-app, email, toast queue across all events.', 'Notifications: ')
img(os.path.join(DIAG,'d7_modules.png'), width=6.4)
caption('Figure 1 — System modules and how they connect.')

# ===== 3. END-TO-END JOURNEY =====
h('3. End-to-End User Journey', 1)
flow('Visitor → Marketing site → Lead form / Register → Email verify → Login → KYC → '
     'Open trading account → Deposit → Trade → Withdraw → Ongoing support / partner growth')
img(os.path.join(DIAG,'d1_journey.png'), width=6.5)
caption('Figure 2 — End-to-end client journey (top row left-to-right, bottom row right-to-left).')
step('Visitor discovers the brand via marketing site or partner referral link.')
step('Submits a lead (demo / contact) or registers a live account.')
step('Verifies email, logs into the client portal.')
step('Completes KYC (identity, address, selfie).')
step('Opens a trading account (Standard / Raw Spread / VIP) with chosen leverage.')
step('Deposits funds via preferred method.')
step('Downloads platform, trades, monitors KPIs on dashboard.')
step('Withdraws profits; raises support tickets as needed.')
step('Optionally becomes an IB to refer and earn rebates.')

# ===== 4. LEAD-TO-CLIENT =====
h('4. Lead-to-Client Conversion Workflow', 1)
flow('Lead captured → Nurtured → Registered → Verified → KYC → Funded → Active Client')
step('Lead captured (contact form, demo request, partner referral, newsletter).')
step('Lead stored with source + UTM + referring IB; status = New.')
step('Automated welcome + nurture notifications; assigned to sales/support if configured.')
step('Lead registers a live account → becomes Pending Client.')
step('Email verified + KYC approved → status = Verified.')
step('First deposit cleared → status = Funded / Active Client (conversion complete).')
para('Conversion KPIs: lead→register, register→KYC, KYC→funded. Each transition emits an event for '
     'analytics and (if referred) credits the originating IB.', italic=True, color=GRAY)
img(os.path.join(DIAG,'h2_conversion.png'), width=6.5)
caption('Figure 3 — Lead-to-client conversion with status at each stage.')

# ===== 5. REGISTRATION =====
h('5. Registration & Onboarding Workflow', 1)
flow('Form (multi-step) → Validation → Create user → Email verify → Login → Onboarding checklist')
step('Step 1 — Personal details (name, email, phone, country).')
step('Step 2 — Account preferences (account type, base currency, leverage).')
step('Step 3 — Credentials + agreements (password, T&C, risk disclosure).')
step('Client-side zod validation on each step; cannot advance until valid.')
step('On submit: create user (status = Pending), send verification email.')
step('Email verified → first login → onboarding checklist (verify KYC, open account, deposit).')
para('Demo flow: shorter form → instant demo account with virtual balance, no KYC required.', italic=True, color=GRAY)
img(os.path.join(DIAG,'h3_registration.png'), width=6.5)
caption('Figure 4 — Registration and onboarding flow.')

# ===== 6. KYC =====
h('6. KYC Verification Workflow', 1)
flow('Not Submitted → Submitted/Pending → Under Review → Approved | Rejected (→ Resubmit)')
step('Client opens KYC; sees 3 steps with current statuses.')
step('Identity Verification — upload ID/passport (front/back).')
step('Proof of Address — upload utility bill / bank statement.')
step('Selfie Verification — upload selfie holding ID.')
step('Each upload sets that step to Pending → Compliance Officer reviews.')
step('Officer approves (Approved) or rejects with reason (Rejected → client resubmits).')
step('All three Approved → overall KYC = Verified → unlocks withdrawals & full funding limits.')
para('Gating: deposits may be allowed pre-KYC up to a limit; withdrawals require full KYC. '
     'Progress bar reflects completed steps; compliance message shown throughout.', italic=True, color=GRAY)
img(os.path.join(DIAG,'h4_kyc.png'), width=6.3)
caption('Figure 5 — KYC verification with compliance review.')

# ===== 7. ACCOUNT CREATION =====
h('7. Trading Account Creation Workflow', 1)
flow('Choose type → Configure → Confirm → Provision account number → Active')
step('Client clicks Open New Account (dashboard) or Choose Account (pricing page).')
step('Selects type: Standard / Raw Spread / VIP.')
step('Configures base currency + leverage (within type limits).')
step('Confirms; system provisions a unique account number with zero balance.')
step('Account appears in dashboard table with status = Active.')
table(['Type','Spreads from','Commission','Leverage','Audience'], [
    ['Standard','0.8 pips','Commission-free','1:500','All traders'],
    ['Raw Spread','0.0 pips','$7 / lot','1:500','Experienced (Popular)'],
    ['VIP','0.0 pips','Custom','1:500','High-volume'],
], widths=[1.2,1.2,1.4,1.0,1.8])

# ===== 8. DASHBOARD =====
h('8. Client Dashboard Workflow', 1)
step('On login, land on Dashboard with "Welcome back, {name}".')
step('Animated KPI widgets: Total Balance, Equity, Free Margin, Margin Level.')
step('Accounts table: Account #, Type, Balance, Equity, Status.')
step('Quick actions: Open New Account, Deposit, Withdraw, complete KYC if pending.')
step('Sidebar navigation to all portal modules; notifications bell + profile menu.')
para('Margin Level = (Equity / Used Margin) x 100. Color-coded thresholds warn on low margin.',
     italic=True, color=GRAY)

# ===== 9. DEPOSIT =====
h('9. Deposit Workflow', 1)
flow('Select method → Enter amount → Confirm → Pending → Processing → Completed | Failed')
step('Funds → Deposit tab; choose method (Bank Transfer, Card, E-Wallet, Crypto).')
step('Enter amount + account to credit; see fees, ETA, min/max.')
step('Confirm → transaction created (status = Pending).')
step('Instant methods (Card/E-Wallet/Crypto) → Processing → Completed quickly; Bank Transfer 1–3 days.')
step('On Completed → account balance credited, toast + notification, history updated.')
step('On Failed/Rejected → reason shown, balance unchanged, client may retry.')
img(os.path.join(DIAG,'h5_deposit.png'), width=6.3)
caption('Figure 6 — Deposit flow from method selection to credited balance.')

# ===== 10. WITHDRAWAL =====
h('10. Withdrawal Workflow', 1)
flow('Request → Validate (KYC + balance) → Pending → Finance review → Approved → Paid | Rejected')
step('Funds → Withdraw; select account + method + amount.')
step('System checks: KYC Verified, sufficient free margin, withdrawal limits, AML rules.')
step('Request created (status = Pending) → balance reserved/held.')
step('Finance Officer reviews; Approved → funds sent → status = Paid/Completed.')
step('Rejected → reason shown, hold released, balance restored.')
para('Anti-fraud: withdrawals typically returned to the original deposit source first '
     '(source-of-funds rule).', italic=True, color=GRAY)
img(os.path.join(DIAG,'h6_withdrawal.png'), width=6.3)
caption('Figure 7 — Withdrawal flow with verification and finance approval.')

# ===== 11. INTERNAL TRANSFER =====
h('11. Internal Transfer Workflow', 1)
flow('Select from/to accounts → Amount → Validate → Execute → Both balances updated')
step('Funds → Internal Transfer; choose source + destination account (same client).')
step('Enter amount; validate sufficient free margin in source.')
step('Confirm → debit source, credit destination (instant, same-currency) or FX-converted.')
step('Two linked transactions recorded; both balances + history updated; toast confirms.')

# ===== 12. SUPPORT =====
h('12. Support / Helpdesk Workflow', 1)
flow('Open ticket → Open → In Progress → Awaiting Reply → Resolved → Closed')
step('Client creates ticket (category, subject, message, optional attachment) or starts live chat.')
step('Ticket queued (status = Open) + reference ID; client + agent notified.')
step('Agent responds → In Progress; threaded replies; status toggles Awaiting Reply.')
step('Issue solved → Resolved; client confirms or auto-closes after inactivity → Closed.')
step('Client can reopen a closed ticket within a window.')

# ===== 13. PARTNER / IB =====
h('13. Partner / IB Workflow', 1)
flow('Apply → Approved → Get referral link → Refer clients → Track volume → Earn rebates → Payout')
step('Visitor applies via Partnership page (Become an IB Partner).')
step('Admin reviews application → Approved → IB account created.')
step('IB receives unique referral link(s) + marketing assets.')
step('Referred leads tagged to IB through registration → conversion.')
step('System tracks referred clients’ trading volume in real time.')
step('Rebates calculated per lot/spread; shown in IB reports dashboard.')
step('IB requests rebate payout → Finance approves → paid; reflected in history.')

# ===== 14. ADMIN =====
h('14. Admin Workflow', 1)
bullet('Create/suspend/edit clients, reset access, adjust leverage.', 'User mgmt: ')
bullet('Review queue, approve/reject KYC with reasons, AML flags.', 'Compliance: ')
bullet('Approve deposits/withdrawals, reconcile, manual adjustments with audit trail.', 'Finance: ')
bullet('Configure account types, spreads, leverage caps, fees.', 'Products: ')
bullet('Approve IBs, set rebate tiers, review partner reports, payouts.', 'Partners: ')
bullet('Edit marketing content, markets list, downloads, announcements.', 'CMS: ')
bullet('Assign tickets, manage agents, SLA monitoring.', 'Support ops: ')
bullet('Immutable audit log of all privileged actions.', 'Audit: ')

# ===== 15. NOTIFICATION =====
h('15. Notification Workflow', 1)
table(['Trigger','Channel','Recipient'], [
    ['Registration / email verify','Email + in-app','Client'],
    ['KYC submitted / approved / rejected','Email + toast','Client + Compliance'],
    ['Deposit / withdrawal status change','Email + toast','Client + Finance'],
    ['Account opened','Toast + in-app','Client'],
    ['Support reply','Email + in-app','Client / Agent'],
    ['IB: new referral / rebate','In-app + email','Partner'],
    ['Security (login, password change)','Email','Client'],
], widths=[2.6,1.7,1.7])
para('Global ToastContext queues transient UI toasts; persistent items live in a notifications center '
     'with read/unread state.', italic=True, color=GRAY)

# ===== 16. COMPLIANCE & RISK =====
h('16. Compliance & Risk Workflow', 1)
step('KYC/AML gate before withdrawals and high-value funding.')
step('Sanctions / PEP screening on registration (integration-ready).')
step('Source-of-funds rule on withdrawals (return to origin).')
step('Transaction monitoring for unusual patterns → risk flag → manual review.')
step('Leverage caps and negative-balance protection per account type/jurisdiction.')
step('Risk disclosure + T&C acceptance recorded at registration.')
step('Full audit trail; data retention + privacy compliance.')

# ===== 17. EDGE CASES =====
h('17. Edge Cases / Exception Handling', 1)
bullet('Allow re-send verification; block login until verified.', 'Unverified email: ')
bullet('Block withdrawal, prompt to complete KYC.', 'KYC incomplete: ')
bullet('Rejected with reason; client resubmits specific step.', 'KYC rejected: ')
bullet('Reject with insufficient-funds message; suggest deposit.', 'Withdrawal > balance: ')
bullet('Show Failed state, keep balance unchanged, offer retry.', 'Payment failure: ')
bullet('Idempotent submit; disable button + spinner; dedupe by request key.', 'Double submit: ')
bullet('Session expiry → redirect to login, preserve intended route.', 'Expired session: ')
bullet('Skeletons on load; ErrorState with retry on fetch failure; EmptyState when no data.', 'Loading/empty/error: ')
bullet('prefers-reduced-motion disables 3D/parallax; static fallbacks.', 'Reduced motion: ')
bullet('Block negative balances; margin-call/stop-out warnings.', 'Negative balance: ')

# ===== 18. STATUSES =====
h('18. Recommended Statuses per Module', 1)
table(['Module','Statuses'], [
    ['Lead','New, Contacted, Qualified, Converted, Lost'],
    ['Client/User','Pending, Verified, Active, Suspended, Closed'],
    ['KYC step','Not Submitted, Pending, Under Review, Approved, Rejected'],
    ['Trading Account','Active, Inactive, Suspended, Archived'],
    ['Deposit','Pending, Processing, Completed, Failed, Cancelled'],
    ['Withdrawal','Pending, Approved, Paid, Rejected, Cancelled'],
    ['Transfer','Pending, Completed, Failed'],
    ['Support Ticket','Open, In Progress, Awaiting Reply, Resolved, Closed'],
    ['IB Application','Pending, Approved, Rejected, Suspended'],
    ['Rebate Payout','Accrued, Requested, Approved, Paid'],
], widths=[1.8,5.0])

# ===== 19. DATA ENTITIES =====
h('19. Database Entities / Data Structure Overview', 1)
table(['Entity','Key fields'], [
    ['User','id, name, email, phone, country, role, status, createdAt'],
    ['Lead','id, source, utm, referrerIB, status, contactInfo'],
    ['KycRecord','userId, idStatus, addressStatus, selfieStatus, docs[], reviewedBy'],
    ['TradingAccount','id, userId, accountNumber, type, currency, leverage, balance, equity, status'],
    ['Transaction','id, accountId, type(deposit/withdraw/transfer), method, amount, status, ref, ts'],
    ['SupportTicket','id, userId, category, subject, messages[], status, agentId'],
    ['Partner/IB','id, userId, status, referralCode, tier, referredClients[], rebateBalance'],
    ['Notification','id, userId, type, message, read, ts'],
    ['AuditLog','id, actorId, action, target, metadata, ts'],
], widths=[1.8,5.0])
para('In the current build these are mock entities in React Context persisted to localStorage; the schema is '
     'backend-ready for a real API.', italic=True, color=GRAY)

# ===== 20. API/MODULE LOGIC =====
h('20. API / Module Logic Overview', 1)
table(['Domain','Representative operations'], [
    ['Auth','register, verifyEmail, login, logout, refreshSession, forgotPassword'],
    ['KYC','submitStep, getStatus, review (admin)'],
    ['Accounts','createAccount, listAccounts, getAccount, updateLeverage'],
    ['Funds','deposit, withdraw, transfer, listTransactions'],
    ['Support','createTicket, replyTicket, listTickets, closeTicket'],
    ['Partner','applyIB, getReferralLink, getRebateReport, requestPayout'],
    ['Admin','manageUsers, reviewKyc, approveFunds, manageContent, viewAudit'],
    ['Notifications','push, list, markRead'],
], widths=[1.6,5.2])
para('Current implementation: context reducers + hooks (useAuth, usePortalData, useToast) standing in for API '
     'calls; same signatures map 1:1 to future REST/GraphQL endpoints.', italic=True, color=GRAY)

# ===== 21. NAVIGATION LOGIC =====
h('21. Page-to-Page Navigation Logic', 1)
bullet('Home → CTAs → Register / Demo / Accounts; nav → About/Markets/Partnership/Contact.', 'Public: ')
bullet('Register → email verify → Login → Dashboard.', 'Onboarding: ')
bullet('Dashboard → KYC (if pending) → Accounts → Funds → trade; sidebar persists across portal.', 'Portal: ')
bullet('Guarded routes: unauthenticated portal access → Login (then redirect back).', 'Guards: ')
bullet('KYC-gated actions (withdraw) → redirect to KYC if incomplete.', 'Gating: ')
bullet('Contextual deep links: notification → relevant page; ticket reply → ticket thread.', 'Deep links: ')

# ===== 22. PERMISSIONS =====
h('22. Suggested Permissions by Role', 1)
table(['Capability','Visitor','Client','IB','Support','Compliance','Finance','Admin'], [
    ['Browse public site','Yes','Yes','Yes','Yes','Yes','Yes','Yes'],
    ['Portal / trade','No','Yes','Yes','-','-','-','Yes'],
    ['Submit KYC','No','Yes','Yes','-','-','-','-'],
    ['Review KYC','No','No','No','No','Yes','No','Yes'],
    ['Approve funds','No','No','No','No','No','Yes','Yes'],
    ['Manage tickets','No','Own','Own','Yes','-','-','Yes'],
    ['Partner reports','No','No','Yes','-','-','-','Yes'],
    ['Manage users/content','No','No','No','No','No','No','Yes'],
], widths=[1.9,0.7,0.7,0.6,0.8,0.95,0.8,0.7])

# ===== 23. FUTURE INTEGRATIONS =====
h('23. Recommended Future Integrations', 1)
bullet('MT4/MT5 or cTrader bridge for real trading + live balances.', 'Trading platform: ')
bullet('Stripe / bank rails / e-wallets / crypto gateway for real money movement.', 'Payments: ')
bullet('Sumsub / Onfido / Jumio for automated identity verification.', 'KYC/AML: ')
bullet('Real-time quotes feed + TradingView charts.', 'Market data: ')
bullet('CRM (HubSpot/Salesforce) + email automation for lead nurture.', 'CRM/marketing: ')
bullet('SendGrid (email), Twilio (SMS/OTP), web push.', 'Notifications: ')
bullet('GA4 / Mixpanel + server-side conversion + affiliate tracking.', 'Analytics: ')
bullet('Authenticator / SMS 2FA, device management.', 'Security: ')

# ===== 24. COSTS, LICENSING & GO-LIVE =====
h('24. Costs, Licensing & Go-Live', 1)
para('Current status:', bold=True)
para('This deliverable is a front-end application using simulated (mock) data persisted locally. Live trading '
     'is NOT active in the current build — no real orders, balances, payments, or identity checks are processed.',
     italic=True, color=GRAY)
para('When live trading starts:', bold=True)
para('Live trading begins only after the following are completed and approved. Each is a separate workstream '
     'beyond the current front-end build:')
bullet('Backend / API layer connecting the app to real services.', 'Backend: ')
bullet('Trading platform / liquidity integration (e.g. MT4 / MT5 / cTrader bridge) for real orders and balances.', 'Trading engine: ')
bullet('Payment gateway onboarding and approval (cards, bank, e-wallet, crypto).', 'Payments: ')
bullet('Automated identity-verification provider connected and live.', 'KYC/AML: ')
bullet('Real-time price/market-data feed subscribed and connected.', 'Market data: ')
bullet('Brokerage regulatory licensing and compliance sign-off in the operating jurisdiction.', 'Licensing: ')
para('Third-party costs (paid services — not free):', bold=True)
para('The integrations in Section 23 are commercial third-party services. Their API keys are paid credentials '
     'with setup and/or ongoing usage charges, billed by each provider and NOT included in this build:')
table(['Service','Typical cost model'], [
    ['Trading platform / bridge','License + monthly/setup fees'],
    ['Payment gateways','Per-transaction % + fixed fee, plus setup'],
    ['KYC / AML provider','Per-verification (per-check) fee'],
    ['Market-data / price feed','Monthly subscription'],
    ['Email / SMS / OTP','Usage-based (per message)'],
    ['Cloud hosting, domain, SSL, monitoring','Recurring subscription'],
], widths=[2.8,4.0])
para('Exact pricing depends on the chosen provider, transaction volume, and region, and will be confirmed in a '
     'separate commercial quotation. API keys/credentials for these services are not free of cost.',
     italic=True, color=GRAY)

# ===== 25. MASTER SUMMARY =====
h('25. Final Master Workflow Summary', 1)
flow('VISITOR → (marketing / IB referral) → REGISTER → VERIFY EMAIL → LOGIN → KYC → '
     'OPEN ACCOUNT → DEPOSIT → TRADE → WITHDRAW → SUPPORT / GROW (IB)')
para('A visitor is captured as a lead and nurtured into a registered user. After email verification and KYC '
     'approval, the client opens a trading account, funds it, and trades — monitored from a KPI dashboard. '
     'Money movement (deposit, withdrawal, transfer) flows through validated, status-driven workflows with '
     'finance and compliance gates. Support, notifications, and an optional IB partner track run throughout, '
     'while admins govern users, compliance, finance, content, and partners behind the scenes. Every module is '
     'status-driven, auditable, permission-scoped, and integration-ready — built today as a mock-data SPA that '
     'maps cleanly onto a real backend.')

# =====================================================================
# APPENDICES — DESIGN & ARCHITECTURE
# =====================================================================
doc.add_page_break()
apx = doc.add_paragraph(); ar = apx.add_run('Appendices — Design & Architecture'); ar.bold=True; ar.font.size=Pt(20); ar.font.color.rgb=RED
doc.add_paragraph()

# ---- Appendix A: Component System ----
h('Appendix A — Reusable Component System', 1)
h('A.1 UI primitives (components/ui)', 2)
para('Button, Card, Badge/StatusChip, Tabs, Modal, Dropdown, Toast, Accordion, Input, Select, FileUpload, '
     'Skeleton, EmptyState, ErrorState, StatCard, ProgressBar, DataTable.')
h('A.2 Layout components', 2)
para('Navbar (responsive + mobile drawer), Footer, PortalSidebar, PortalTopbar, PublicLayout, PortalLayout, '
     'Container, Section.')
h('A.3 Marketing / composite components', 2)
para('HeroBlock, StatsStrip, FeatureCard, MarketCategoryCard, AccountComparisonCard, PartnerBenefit, '
     'InfinityRibbon, ContactForm, AnimatedCounter, ScrollReveal.')
h('A.4 Portal components', 2)
para('KpiWidget, AccountsTable, DepositMethodCard, WithdrawForm, TransferForm, TransactionHistory, '
     'KycStepCard, SupportTicketForm, DownloadCard, RegisterStepper, LoginForm.')

# ---- Appendix B: Folder Structure ----
h('Appendix B — Recommended Folder Structure', 1)
mono = doc.add_paragraph()
mr = mono.add_run(
"""src/
  main.tsx
  App.tsx
  routes.tsx
  assets/                 # logos, textures, fonts
  styles/                 # globals.css, tailwind layers
  lib/                    # cn(), formatters, constants
  data/                   # mock accounts, markets, transactions
  context/                # AuthContext, PortalDataContext, ToastContext
  hooks/                  # useAuth, usePortalData, useToast, useReducedMotion
  animations/             # framer variants, gsap timelines
  components/
    ui/                   # Button, Card, Tabs, Modal, ...
    layout/               # Navbar, Footer, Sidebar, layouts
    three/                # HeroScene, Globe, InfinityRibbon, Particles
    marketing/            # HeroBlock, StatsStrip, FeatureCard, ...
    portal/               # KpiWidget, AccountsTable, DepositMethodCard, ...
  pages/
    public/               # Home, About, Accounts, Markets, Partnership, Contact
    auth/                 # Login, Register, Demo
    portal/               # Dashboard, Accounts, Funds, Kyc, Downloads, Profile, Support
""")
mr.font.name='Consolas'; mr.font.size=Pt(9)

# ---- Appendix C: 3D / Animation Plan ----
h('Appendix C — 3D & Animation Plan', 1)
table(['Element','Tech','Behavior'], [
    ['Hero laptop + mobile + waveform','R3F / Three.js','Floating mockups over glowing red market wave + particles'],
    ['Rotating globe + arcs','R3F / Three.js','Slow auto-rotate, glowing connection arcs (partner section)'],
    ['Infinity ribbon','R3F / shader','Glowing red looping light ribbon (partnership page)'],
    ['Section backgrounds','GSAP / canvas','Subtle red particle field + animated waveform'],
    ['Scroll reveals / parallax','Framer Motion','Layered fade/slide on enter, parallax depth'],
    ['KPI counters','Framer Motion','Count-up on in-view'],
    ['Cards / buttons','Framer + CSS','Lift + red edge-glow on hover'],
    ['Nav / tabs','Framer Motion','Animated active indicator, layout transitions'],
    ['KYC / funds progress','Framer Motion','Animated progress + status transitions'],
], widths=[2.3,1.3,2.8])
para('Accessibility: all 3D is lazy-loaded behind Suspense + skeleton; prefers-reduced-motion renders static '
     'fallbacks. Easing is fast and elegant — premium, never excessive.', italic=True, color=GRAY)

# ---- Appendix D: Design System ----
h('Appendix D — Design System', 1)
h('D.1 Color palette', 2)
table(['Token','Value','Use'], [
    ['bg-base','#050505','App background (near-black)'],
    ['bg-elevated','#0a0a0a','Elevated background'],
    ['panel','#111111 -> #181818','Cards / glass panels'],
    ['primary (red)','#e11d2e','Accent, CTAs, glows, active states'],
    ['text','#ffffff','Primary text'],
    ['text-muted','#9aa0a6','Secondary text'],
    ['border','#222222','Thin dark borders (red on active)'],
    ['success','#16a34a','Status: Active / Approved'],
    ['warning','#f59e0b','Status: Pending'],
    ['danger','#dc2626','Status: Rejected'],
], widths=[1.6,1.9,2.9])
h('D.2 Typography', 2)
bullet('Inter (400-800) - body, UI, data.', 'Primary: ')
bullet('Space Grotesk (500-700) - display headlines, brand.', 'Display: ')
bullet('Scale: 12 / 14 / 16 / 20 / 24 / 32 / 48 / 64px with tight display leading.', 'Scale: ')
h('D.3 Spacing & utilities', 2)
bullet('4px base scale (4/8/12/16/24/32/48/64/96).', 'Spacing: ')
bullet('.glass-panel, .red-glow, .card-lift; consistent radii; reduced-motion fallbacks everywhere.', 'Utilities: ')

# ---- Appendix E: Build Phases ----
h('Appendix E — Implementation Phases', 1)
table(['Phase','Scope','Deliverables'], [
    ['1 - Foundation','Scaffold + design system','Tailwind tokens, UI primitives, contexts, mock data, toasts, layouts'],
    ['2 - Marketing','Public website','Navbar/Footer, Home + all public pages, 3D hero + globe, scroll reveals'],
    ['3 - Auth & Shell','Login + portal shell','Auth flow, route guards, sidebar shell, dashboard + KPIs + accounts table'],
    ['4 - Portal modules','Feature depth','Funds, KYC, downloads, profile, support, multi-step onboarding'],
    ['5 - Polish & QA','Production readiness','Skeleton/empty/error states, responsive, a11y, reduced-motion, build QA'],
], widths=[1.5,1.6,3.3])

if os.path.exists(LOGO):
    make_watermark(LOGO, WM)
    set_watermark(doc.sections[0], WM)

out = r'c:/Users/gowth/OneDrive/Desktop/27 trading/27-Markets-Complete-Workflow.docx'
try:
    doc.save(out)
    print('Saved:', out)
except PermissionError:
    alt = r'c:/Users/gowth/OneDrive/Desktop/27 trading/27-Markets-Complete-Workflow-v2.docx'
    doc.save(alt)
    print('Original file is open/locked. Saved to:', alt)
