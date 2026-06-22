# -*- coding: utf-8 -*-
"""Client-facing 27 Markets CRM (Back-Office) overview — what it includes and why.
Matches the branding/engine of build_client_doc.py."""
import os
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch, Polygon
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image

BASE = r'c:/Users/gowth/OneDrive/Desktop/27 trading'
DIAG = os.path.join(BASE, 'assets', 'diagrams')
os.makedirs(DIAG, exist_ok=True)
LOGO = r'C:/Users/gowth/Downloads/corecare consultancy.png'
WM = os.path.join(BASE, 'assets', 'watermark.png')

def make_watermark(src, dst, opacity=0.10):
    im = Image.open(src).convert('RGBA')
    bg = Image.new('RGBA', im.size, (255, 255, 255, 255))
    comp = Image.alpha_composite(bg, im)
    Image.blend(bg, comp, opacity).convert('RGB').save(dst)

def set_watermark(section, image_path, width_in=4.6):
    header = section.header
    header.is_linked_to_previous = False
    run = header.add_paragraph().add_run()
    run.add_picture(image_path, width=Inches(width_in))
    drawing = run._r.find(qn('w:drawing'))
    inline = drawing.find(qn('wp:inline'))
    extent = inline.find(qn('wp:extent'))
    cx, cy = extent.get('cx'), extent.get('cy')
    docPr = inline.find(qn('wp:docPr'))
    graphic = inline.find(qn('a:graphic'))
    anchor = OxmlElement('wp:anchor')
    for k, v in {'distT':'0','distB':'0','distL':'0','distR':'0','simplePos':'0',
                 'relativeHeight':'251650048','behindDoc':'1','locked':'0',
                 'layoutInCell':'1','allowOverlap':'1'}.items():
        anchor.set(k, v)
    sp = OxmlElement('wp:simplePos'); sp.set('x','0'); sp.set('y','0'); anchor.append(sp)
    ph = OxmlElement('wp:positionH'); ph.set('relativeFrom','page')
    a1 = OxmlElement('wp:align'); a1.text='center'; ph.append(a1); anchor.append(ph)
    pv = OxmlElement('wp:positionV'); pv.set('relativeFrom','page')
    a2 = OxmlElement('wp:align'); a2.text='center'; pv.append(a2); anchor.append(pv)
    ex = OxmlElement('wp:extent'); ex.set('cx',cx); ex.set('cy',cy); anchor.append(ex)
    ee = OxmlElement('wp:effectExtent')
    for k in ('l','t','r','b'): ee.set(k,'0')
    anchor.append(ee)
    anchor.append(OxmlElement('wp:wrapNone'))
    anchor.append(docPr)
    anchor.append(OxmlElement('wp:cNvGraphicFramePr'))
    anchor.append(graphic)
    drawing.remove(inline); drawing.append(anchor)

# ---- brand colors ----
RED      = '#e11d2e'
CHARCOAL = '#1b1d22'
PANEL    = '#2a2d34'
INK      = '#111111'
PAGE_BG  = '#ffffff'
RED_RGB  = RGBColor(0xE1, 0x1D, 0x2E)
GRAY_RGB = RGBColor(0x55, 0x55, 0x55)
WHITE_RGB= RGBColor(0xFF, 0xFF, 0xFF)

# =====================================================================
# DIAGRAM ENGINE  (shared with build_client_doc.py)
# =====================================================================
def _anchor(cx, cy, w, h, tx, ty):
    dx, dy = tx - cx, ty - cy
    if abs(dy) >= abs(dx):
        return (cx, cy - h/2) if dy < 0 else (cx, cy + h/2)
    return (cx - w/2, cy) if dx < 0 else (cx + w/2, cy)

def draw_flow(nodes, edges, path, figsize, title=None):
    fig, ax = plt.subplots(figsize=figsize, dpi=200)
    ax.set_facecolor(PAGE_BG); fig.patch.set_facecolor(PAGE_BG)
    for e in edges:
        a, b = e[0], e[1]
        label = e[2] if len(e) > 2 else None
        ax_, ay_, aw, ah = nodes[a][0], nodes[a][1], nodes[a][2], nodes[a][3]
        bx_, by_, bw, bh = nodes[b][0], nodes[b][1], nodes[b][2], nodes[b][3]
        p1 = _anchor(ax_, ay_, aw, ah, bx_, by_)
        p2 = _anchor(bx_, by_, bw, bh, ax_, ay_)
        arr = FancyArrowPatch(p1, p2, arrowstyle='-|>', mutation_scale=14,
                              lw=1.6, color=RED, shrinkA=0, shrinkB=0,
                              connectionstyle='arc3,rad=0.0', zorder=1)
        ax.add_patch(arr)
        if label:
            mx, my = (p1[0]+p2[0])/2, (p1[1]+p2[1])/2
            ax.text(mx, my, label, fontsize=7.5, color=INK, ha='center', va='center',
                    bbox=dict(boxstyle='round,pad=0.18', fc='white', ec='#cccccc', lw=0.6), zorder=3)
    for name, (cx, cy, w, h, label, kind) in nodes.items():
        if kind == 'decision':
            pts = [(cx, cy+h/2), (cx+w/2, cy), (cx, cy-h/2), (cx-w/2, cy)]
            poly = Polygon(pts, closed=True, fc='white', ec=RED, lw=1.8, zorder=2)
            ax.add_patch(poly)
            ax.text(cx, cy, label, ha='center', va='center', fontsize=7.6, color=INK, zorder=3)
        elif kind == 'terminal':
            box = FancyBboxPatch((cx-w/2, cy-h/2), w, h, boxstyle='round,pad=0.02,rounding_size=0.18',
                                 fc=RED, ec=RED, lw=1.5, zorder=2)
            ax.add_patch(box)
            ax.text(cx, cy, label, ha='center', va='center', fontsize=8.2, color='white',
                    fontweight='bold', zorder=3)
        elif kind == 'group':
            box = FancyBboxPatch((cx-w/2, cy-h/2), w, h, boxstyle='round,pad=0.02,rounding_size=0.10',
                                 fc='none', ec=RED, lw=1.4, ls='--', zorder=2)
            ax.add_patch(box)
            ax.text(cx, cy+h/2-0.16, label, ha='center', va='top', fontsize=8.4, color=RED,
                    fontweight='bold', zorder=3)
        else:  # process
            box = FancyBboxPatch((cx-w/2, cy-h/2), w, h, boxstyle='round,pad=0.02,rounding_size=0.12',
                                 fc=CHARCOAL, ec=PANEL, lw=1.2, zorder=2)
            ax.add_patch(box)
            ax.text(cx, cy, label, ha='center', va='center', fontsize=8.0, color='white', zorder=3)
    xs = [v[0] for v in nodes.values()]; ys = [v[1] for v in nodes.values()]
    ws = [v[2] for v in nodes.values()]; hs = [v[3] for v in nodes.values()]
    ax.set_xlim(min(x-w/2 for x,w in zip(xs,ws))-0.4, max(x+w/2 for x,w in zip(xs,ws))+0.4)
    ax.set_ylim(min(y-h/2 for y,h in zip(ys,hs))-0.4, max(y+h/2 for y,h in zip(ys,hs))+0.4)
    ax.set_aspect('equal'); ax.axis('off')
    if title:
        ax.set_title(title, fontsize=11, color=INK, fontweight='bold', pad=10)
    fig.savefig(path, bbox_inches='tight', facecolor=PAGE_BG, dpi=200)
    plt.close(fig)
    return path

def P(cx,cy,label,w=2.6,h=0.9,kind='process'): return (cx,cy,w,h,label,kind)

# =====================================================================
# BUILD DIAGRAMS
# =====================================================================
# --- A. CRM back-office module map ---
n = {
 'dash':(-3.5, 1.7, 3.0, 1.2, 'DASHBOARD\nKPIs · Activity\nAlerts','group'),
 'cli':(0, 1.7, 3.0, 1.2, 'CLIENTS\n360 View · Notes\nAccounts · History','group'),
 'lead':(3.5, 1.7, 3.0, 1.2, 'LEADS\nPipeline · Assign\nFollow-ups','group'),
 'kyc':(-3.5, -0.7, 3.0, 1.2, 'KYC / COMPLIANCE\nReview Queue\nApprove · Reject','group'),
 'fin':(0, -0.7, 3.0, 1.2, 'FINANCE\nDeposits · Withdrawal\nApprovals','group'),
 'sup':(3.5, -0.7, 3.0, 1.2, 'SUPPORT\nTicket Inbox\nAssign · Reply','group'),
 'ib':(-3.5, -3.1, 3.0, 1.2, 'PARTNERS / IB\nReferrals\nRebates','group'),
 'rep':(0, -3.1, 3.0, 1.2, 'REPORTS\nRevenue · Funnel\nPerformance','group'),
 'staff':(3.5, -3.1, 3.0, 1.2, 'STAFF & SETTINGS\nRoles · Audit Log\nAccess Control','group'),
}
edges = [('dash','cli'),('cli','lead'),('cli','kyc'),('cli','fin'),
         ('cli','sup'),('cli','ib'),('dash','rep'),('staff','fin')]
draw_flow(n, edges, os.path.join(DIAG,'crm_modules.png'), (8.6, 6.0))

# --- B. Lead → funded client pipeline (horizontal) ---
GX, MW, MH, BRY = 3.0, 2.6, 1.15, -2.35
def hnode(i, label, kind='process', row=0):
    return P(i*GX, 0 if row==0 else BRY, label, w=MW, h=MH, kind=kind)
def hdiamond(i, label, row=0):
    return (i*GX, 0 if row==0 else BRY, 2.7, 1.35, label, 'decision')
labels=['New\nLead','Assigned\nto Agent','Contacted\n/ Qualified','Registered\n+ Verified','Funded','Active\nClient']
n={f'c{i}':hnode(i,l,'terminal' if i in (0,5) else 'process') for i,l in enumerate(labels)}
edges=[(f'c{i}',f'c{i+1}') for i in range(5)]
draw_flow(n, edges, os.path.join(DIAG,'crm_pipeline.png'), (15.5, 2.4))

# --- C. KYC review (back-office, with reject loop) ---
n={'a':hnode(0,'Client\nSubmits','terminal'),'b':hnode(1,'Review\nQueue'),
   'c':hnode(2,'Officer\nChecks Docs'),'d':hdiamond(3,'Compliance\nApproves?'),
   'e':hnode(4,'Verified\n(Unlocks)','terminal'),
   'r':hnode(3,'Rejected:\nResubmit',row=1)}
edges=[('a','b'),('b','c'),('c','d'),('d','e','Yes'),('d','r','No')]
draw_flow(n, edges, os.path.join(DIAG,'crm_kyc.png'), (13.5, 4.3))

# --- D. Withdrawal approval (finance back-office) ---
n={'a':hnode(0,'Client\nRequests','terminal'),'b':hdiamond(1,'KYC + Funds\nValid?'),
   'c':hnode(2,'Pending\n(Held)'),'d':hdiamond(3,'Finance\nApproves?'),
   'e':hnode(4,'Paid /\nCompleted','terminal'),
   'b2':hnode(1,'Blocked:\nComplete KYC',row=1),'e2':hnode(3,'Rejected:\nHold Released',row=1)}
edges=[('a','b'),('b','c','Yes'),('b','b2','No'),('c','d'),('d','e','Yes'),('d','e2','No')]
draw_flow(n, edges, os.path.join(DIAG,'crm_withdrawal.png'), (13.5, 4.3))

# --- E. Two-tier access ---
n={
 'adm':(0, 0, 3.0, 1.2, 'ADMIN\nFull access — finance\napprovals, roles,\nsettings, all modules','group'),
 'agt':(3.8, 0, 3.0, 1.2, 'AGENT\nClients · Leads\nSupport · KYC review\n(no finance / roles)','group'),
 'aud':(1.9, -2.3, 3.4, 1.0, 'AUDIT LOG\nEvery action recorded','terminal'),
}
edges=[('adm','aud'),('agt','aud')]
draw_flow(n, edges, os.path.join(DIAG,'crm_roles.png'), (8.0, 4.0))

print('Diagrams rendered to', DIAG)

# =====================================================================
# BUILD DOCX
# =====================================================================
doc = Document()
nm = doc.styles['Normal']; nm.font.name='Calibri'; nm.font.size=Pt(10.5)
nm.paragraph_format.space_after=Pt(6)
for s in doc.sections:
    s.top_margin=Inches(0.8); s.bottom_margin=Inches(0.8)
    s.left_margin=Inches(0.9); s.right_margin=Inches(0.9)

def shade(cell,hexc):
    tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),hexc); tcPr.append(shd)
def H(text,level=1):
    p=doc.add_heading(text,level=level)
    for r in p.runs: r.font.color.rgb = RED_RGB if level==1 else RGBColor(0x11,0x11,0x11)
    return p
def para(text,bold=False,italic=False,color=None,size=None,after=6):
    p=doc.add_paragraph(); r=p.add_run(text); r.bold=bold; r.italic=italic
    if color:r.font.color.rgb=color
    if size:r.font.size=Pt(size)
    p.paragraph_format.space_after=Pt(after); return p
def bullet(text,prefix=None):
    p=doc.add_paragraph(style='List Bullet')
    if prefix: rr=p.add_run(prefix); rr.bold=True; p.add_run(text)
    else: p.add_run(text)
    return p
def img(path,width=6.3):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Inches(width)); return p
def caption(text):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(text); r.italic=True; r.font.size=Pt(8.5); r.font.color.rgb=GRAY_RGB; return p
def table(headers,rows,widths=None):
    t=doc.add_table(rows=1,cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style='Light Grid Accent 1'
    for i,head in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=''; run=c.paragraphs[0].add_run(head)
        run.bold=True; run.font.color.rgb=WHITE_RGB; run.font.size=Pt(9); shade(c,'E11D2E')
    for row in rows:
        cells=t.add_row().cells
        for i,v in enumerate(row):
            cells[i].text=''; run=cells[i].paragraphs[0].add_run(str(v)); run.font.size=Pt(9)
    if widths:
        for i,w in enumerate(widths):
            for r_ in t.rows: r_.cells[i].width=Inches(w)
    doc.add_paragraph(); return t
def set_valign(cell, val='center'):
    tcPr=cell._tc.get_or_add_tcPr(); v=OxmlElement('w:vAlign'); v.set(qn('w:val'),val); tcPr.append(v)
def no_split(row):
    trPr=row._tr.get_or_add_trPr(); cs=OxmlElement('w:cantSplit'); trPr.append(cs)
def side_by_side(image_path, img_w, lines, img_left=True):
    t=doc.add_table(rows=1, cols=2); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False
    no_split(t.rows[0])
    ci, ct = (0,1) if img_left else (1,0)
    cimg=t.rows[0].cells[ci]; ctxt=t.rows[0].cells[ct]
    cimg.width=Inches(3.7); ctxt.width=Inches(2.8)
    set_valign(cimg); set_valign(ctxt)
    p=cimg.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(image_path, width=Inches(img_w))
    first=True
    for ln in lines:
        p = ctxt.paragraphs[0] if first else ctxt.add_paragraph(); first=False
        rb=p.add_run('▸ '); rb.font.color.rgb=RED_RGB; rb.bold=True
        p.add_run(ln).font.size=Pt(9.5)
        p.paragraph_format.space_after=Pt(5)
    doc.add_paragraph(); return t
def field(p,code):
    for t,attr in [('begin',None),('inst',code),('separate',None),('text','?'),('end',None)]:
        run=p.add_run()
        if t=='inst':
            el=OxmlElement('w:instrText'); el.set(qn('xml:space'),'preserve'); el.text=attr; run._r.append(el)
        elif t=='text': run.text=attr
        else:
            el=OxmlElement('w:fldChar'); el.set(qn('w:fldCharType'),t); run._r.append(el)

# footer page numbers
fp=doc.sections[0].footer.paragraphs[0]; fp.alignment=WD_ALIGN_PARAGRAPH.CENTER
fp.add_run('27 Markets — CRM (Back-Office) Overview   |   Page ').font.size=Pt(8)
field(fp,'PAGE'); fp.add_run(' of ').font.size=Pt(8); field(fp,'NUMPAGES')

# ---- COVER ----
t=doc.add_paragraph(); t.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=t.add_run('27 MARKETS'); r.bold=True; r.font.size=Pt(36); r.font.color.rgb=RED_RGB
s=doc.add_paragraph(); s.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=s.add_run('Trade Beyond Limits'); r.italic=True; r.font.size=Pt(13); r.font.color.rgb=GRAY_RGB
doc.add_paragraph()
s2=doc.add_paragraph(); s2.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=s2.add_run('Client Relationship Management (CRM)'); r.bold=True; r.font.size=Pt(20)
s3=doc.add_paragraph(); s3.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=s3.add_run('Back-Office Platform — What It Includes & Why We Are Building It'); r.font.size=Pt(13); r.font.color.rgb=GRAY_RGB
doc.add_paragraph()
m=doc.add_paragraph(); m.alignment=WD_ALIGN_PARAGRAPH.CENTER
m.add_run('Prepared for: Client   |   Date: 2026-06-22   |   Version 1.0').font.size=Pt(10)
doc.add_paragraph()
i=doc.add_paragraph(); i.alignment=WD_ALIGN_PARAGRAPH.CENTER
ir=i.add_run('This document explains the internal back-office system (CRM) that powers the 27 Markets '
             'platform behind the scenes — the tools your team uses to manage leads, clients, identity '
             'verification, deposits and withdrawals, support, and partners — and the business reasons for '
             'building it.')
ir.italic=True; ir.font.color.rgb=GRAY_RGB; ir.font.size=Pt(10.5)
doc.add_page_break()

# ---- TOC ----
th=doc.add_paragraph(); r=th.add_run('Contents'); r.bold=True; r.font.size=Pt(16); r.font.color.rgb=RED_RGB
tp=doc.add_paragraph(); field(tp,'TOC \\o "1-1" \\h \\z \\u')
doc.add_page_break()

# ---- 1. WHAT IS THE CRM ----
H('1. What Is the CRM?',1)
para('The 27 Markets platform has two sides. Clients see the public website and the secure client portal. '
     'Your team needs the other side: a private, secure back-office to run the business. That back-office is '
     'the CRM (Client Relationship Management) system.')
para('In one place, the CRM lets your staff see every lead and client, review identity documents, approve '
     'deposits and withdrawals, answer support requests, manage partners, and read live performance reports — '
     'with strict access controls and a full record of who did what.')
img(os.path.join(DIAG,'crm_modules.png'), width=6.5)
caption('Figure 1 — The CRM back-office modules and how they connect around the client.')

# ---- 2. WHY ----
H('2. Why We Are Building It',1)
para('Without a CRM, client information lives in scattered spreadsheets, inboxes, and chat messages. That is '
     'slow, error-prone, hard to audit, and impossible to scale. The CRM solves this directly:')
bullet('Every lead, client, document, payment, and conversation in one place — no more scattered spreadsheets or lost information.', 'One source of truth — ')
bullet('Leads are captured, assigned, and followed up systematically, so more visitors become funded, active clients.', 'More conversions, more revenue — ')
bullet('A KYC review queue, sanction checks, and a complete audit trail keep the business compliant and inspection-ready.', 'Regulatory compliance — ')
bullet('No money moves without staff review. Withdrawals are held and approved by finance, protecting against fraud and error.', 'Financial control & safety — ')
bullet('Live dashboards show deposits, withdrawals, pending work, and team performance at a glance.', 'Operational visibility — ')
bullet('Staff see a client’s full history instantly, so they resolve requests faster and more personally.', 'Better client service — ')
bullet('Role-based access (Admin / Agent) plus an audit log mean each person sees only what they should, and every action is recorded.', 'Accountability & security — ')
bullet('As the client base grows from hundreds to thousands, the same system keeps working — manual methods would break.', 'Built to scale — ')

# ---- 3. MODULES OVERVIEW ----
H('3. What the CRM Includes',1)
para('The CRM is delivered as a single back-office application with the following modules:')
table(['Module','What your team can do'], [
    ['Dashboard','See key numbers and pending work — new clients, KYC to review, withdrawals to approve, open tickets, deposits today.'],
    ['Clients (360 View)','Search any client and see their full profile, accounts, balances, KYC status, transactions, notes, and activity history.'],
    ['Leads','A visual sales pipeline: capture leads from the website and demo requests, assign them to agents, and track follow-ups to conversion.'],
    ['KYC / Compliance','Review submitted identity documents in a queue and approve or reject them, which unlocks the client’s withdrawals.'],
    ['Finance','See all deposits and approve or reject withdrawal requests; make controlled manual adjustments (Admin only).'],
    ['Accounts','View all trading accounts and suspend, activate, or change leverage when needed.'],
    ['Support','A shared ticket inbox — assign tickets to staff, reply to clients, and keep internal notes.'],
    ['Partners / IB','See partners and the clients they referred, with referral and rebate tracking.'],
    ['Reports','Charts for deposits, withdrawals, net flow, lead conversion, and agent performance.'],
    ['Staff & Settings','Manage staff accounts, assign Admin or Agent access, and review the full audit log (Admin only).'],
], widths=[1.6,4.9])

# ---- 4. LEADS ----
H('4. Lead & Sales Pipeline',1)
para('Every visitor is a potential client. The CRM captures leads automatically from the website and demo '
     'requests, assigns them to a sales agent, and tracks each one through clear stages until they fund and '
     'start trading — so no opportunity is missed.')
img(os.path.join(DIAG,'crm_pipeline.png'), width=6.6)
caption('Figure 2 — How a lead moves to a funded, active client.')

# ---- 5. CLIENT 360 ----
H('5. The Client 360° View',1)
para('When a staff member opens a client, they see everything about that person on one screen — no switching '
     'between tools:')
bullet('Personal profile and contact details.')
bullet('All trading accounts with live balances.')
bullet('KYC verification status and uploaded documents.')
bullet('Full deposit, withdrawal, and transfer history.')
bullet('Internal staff notes and the complete activity timeline.')

# ---- 6. KYC ----
H('6. KYC & Compliance Review',1)
para('To keep the platform secure and compliant, clients verify their identity. Submissions land in a review '
     'queue where a compliance officer approves or rejects them. Withdrawals stay locked until verification is '
     'approved.')
side_by_side(os.path.join(DIAG,'crm_kyc.png'), 4.6, [
    'The client uploads identity, address, and selfie documents.',
    'Submissions appear in the compliance review queue.',
    'An officer checks the documents against the rules.',
    'Approved clients are Verified and withdrawals unlock.',
    'Rejected submissions are returned with a reason to resubmit.',
])
caption('Figure 3 — KYC review handled inside the CRM.')

# ---- 7. FINANCE ----
H('7. Finance & Withdrawal Approvals',1)
para('Money safety is central. Every withdrawal is checked and held, then approved by the finance team before '
     'any payout — so no funds leave without a human review and a recorded decision.')
side_by_side(os.path.join(DIAG,'crm_withdrawal.png'), 4.6, [
    'A client requests a withdrawal in the portal.',
    'The system checks KYC status and available balance.',
    'Valid requests are held and placed in the approval queue.',
    'Finance reviews and approves the payout, or rejects it.',
    'Rejections release the hold and restore the balance.',
], img_left=False)
caption('Figure 4 — Withdrawal approval workflow in the finance module.')

# ---- 8. ROLES ----
H('8. Staff Roles & Access Control',1)
para('Access is two-tier, so each person only sees and does what their job requires — and every action is '
     'logged for accountability.')
side_by_side(os.path.join(DIAG,'crm_roles.png'), 3.6, [
    'Admin: full access — finance approvals, staff and role management, settings, and all modules.',
    'Agent: day-to-day work — clients, leads, support, and KYC review — but no finance approvals or role changes.',
    'Audit log: every state-changing action is recorded with who, what, and when.',
])
caption('Figure 5 — Two-tier access (Admin / Agent) with a full audit trail.')

# ---- 9. STATUS REFERENCE ----
H('9. Status Reference',1)
para('Every item the team manages has a clear status, so work is never ambiguous.')
table(['Area','Possible statuses'], [
    ['Lead','New, Contacted, Qualified, Converted, Lost'],
    ['Client account','Active, Suspended, Closed'],
    ['KYC step','Not Submitted, Pending, Approved, Rejected'],
    ['Trading account','Active, Pending, Suspended, Archived'],
    ['Deposit','Pending, Completed, Failed'],
    ['Withdrawal','Pending, Approved, Paid, Rejected'],
    ['Support ticket','Open, In Progress, Resolved'],
], widths=[1.9,4.6])

# ---- 10. ROLLOUT ----
H('10. How We Will Build It',1)
para('The CRM is delivered in independent phases, so you can see and use working pieces early rather than '
     'waiting for everything at once:')
table(['Phase','Delivered'], [
    ['1','Foundations — staff roles (Admin/Agent), secure access, and the admin shell.'],
    ['2','Dashboard with live numbers and recent activity.'],
    ['3','Clients (360 view + notes) and the KYC review queue.'],
    ['4','Finance (withdrawal approvals) and account administration.'],
    ['5','Leads pipeline and the support desk.'],
    ['6','Partners, reports, staff/settings, audit log, and final polish.'],
], widths=[0.9,5.6])

# ---- 11. COSTS & GO-LIVE ----
H('11. Important Note — Costs & Go-Live',1)
para('Please note the following:', bold=True)
bullet('This build runs on simulated data for demonstration. Live money movement is not active yet.', 'Current stage: ')
bullet('The CRM goes fully live once the trading platform, payment gateways, identity-verification provider, and brokerage licensing are connected and approved.', 'Going live: ')
bullet('Some capabilities (identity-verification checks, email/SMS notifications, payment reconciliation) rely on paid third-party services billed by each provider; they are not free or included in this build.', 'Third-party costs: ')
bullet('Advanced features — automated marketing campaigns, a full IB commission engine, and granular per-permission roles — are planned for later phases and quoted separately.', 'Future phases: ')

# ---- 12. SUMMARY ----
H('12. Summary',1)
para('The CRM is the control center of 27 Markets. It turns scattered, manual work into one secure system '
     'where leads convert faster, clients are served better, money is moved only with proper approval, and the '
     'business stays compliant and fully auditable. It gives your team complete operational visibility today, '
     'and a foundation that scales as 27 Markets grows.')

if os.path.exists(LOGO):
    make_watermark(LOGO, WM)
    set_watermark(doc.sections[0], WM)

out = os.path.join(BASE, '27-Markets-CRM-Overview.docx')
try:
    doc.save(out)
    print('Saved:', out)
except PermissionError:
    alt = os.path.join(BASE, '27-Markets-CRM-Overview-v2.docx')
    doc.save(alt)
    print('Original file is open/locked. Saved to:', alt)
