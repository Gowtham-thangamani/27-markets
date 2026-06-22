# -*- coding: utf-8 -*-
"""Client-facing 27 Markets workflow document with custom-drawn flow diagrams."""
import os
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch, Polygon
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image

BASE = r'c:/Users/gowth/OneDrive/Desktop/27 trading'
DIAG = os.path.join(BASE, 'assets', 'diagrams')
os.makedirs(DIAG, exist_ok=True)
LOGO = r'C:/Users/gowth/Downloads/corecare consultancy.png'
WM = os.path.join(BASE, 'assets', 'watermark.png')

def make_watermark(src, dst, opacity=0.10):
    im = Image.open(src).convert('RGBA')
    bg = Image.new('RGBA', im.size, (255, 255, 255, 255))
    comp = Image.alpha_composite(bg, im)
    Image.blend(bg, comp, opacity).convert('RGB').save(dst)

def set_watermark(section, image_path, width_in=4.6):
    header = section.header
    header.is_linked_to_previous = False
    run = header.add_paragraph().add_run()
    run.add_picture(image_path, width=Inches(width_in))
    drawing = run._r.find(qn('w:drawing'))
    inline = drawing.find(qn('wp:inline'))
    extent = inline.find(qn('wp:extent'))
    cx, cy = extent.get('cx'), extent.get('cy')
    docPr = inline.find(qn('wp:docPr'))
    graphic = inline.find(qn('a:graphic'))
    anchor = OxmlElement('wp:anchor')
    for k, v in {'distT':'0','distB':'0','distL':'0','distR':'0','simplePos':'0',
                 'relativeHeight':'251650048','behindDoc':'1','locked':'0',
                 'layoutInCell':'1','allowOverlap':'1'}.items():
        anchor.set(k, v)
    sp = OxmlElement('wp:simplePos'); sp.set('x','0'); sp.set('y','0'); anchor.append(sp)
    ph = OxmlElement('wp:positionH'); ph.set('relativeFrom','page')
    a1 = OxmlElement('wp:align'); a1.text='center'; ph.append(a1); anchor.append(ph)
    pv = OxmlElement('wp:positionV'); pv.set('relativeFrom','page')
    a2 = OxmlElement('wp:align'); a2.text='center'; pv.append(a2); anchor.append(pv)
    ex = OxmlElement('wp:extent'); ex.set('cx',cx); ex.set('cy',cy); anchor.append(ex)
    ee = OxmlElement('wp:effectExtent')
    for k in ('l','t','r','b'): ee.set(k,'0')
    anchor.append(ee)
    anchor.append(OxmlElement('wp:wrapNone'))
    anchor.append(docPr)
    anchor.append(OxmlElement('wp:cNvGraphicFramePr'))
    anchor.append(graphic)
    drawing.remove(inline); drawing.append(anchor)

# ---- brand colors ----
RED      = '#e11d2e'
CHARCOAL = '#1b1d22'
PANEL    = '#2a2d34'
INK      = '#111111'
PAGE_BG  = '#ffffff'
RED_RGB  = RGBColor(0xE1, 0x1D, 0x2E)
GRAY_RGB = RGBColor(0x55, 0x55, 0x55)
WHITE_RGB= RGBColor(0xFF, 0xFF, 0xFF)

# =====================================================================
# DIAGRAM ENGINE
# =====================================================================
def _anchor(cx, cy, w, h, tx, ty):
    """Pick an edge point on a box facing target (tx,ty)."""
    dx, dy = tx - cx, ty - cy
    if abs(dy) >= abs(dx):
        return (cx, cy - h/2) if dy < 0 else (cx, cy + h/2)
    return (cx - w/2, cy) if dx < 0 else (cx + w/2, cy)

def draw_flow(nodes, edges, path, figsize, title=None):
    """nodes: name -> (cx, cy, w, h, label, kind).  kind: terminal|process|decision|group"""
    fig, ax = plt.subplots(figsize=figsize, dpi=200)
    ax.set_facecolor(PAGE_BG); fig.patch.set_facecolor(PAGE_BG)

    # edges first (under nodes)
    for e in edges:
        a, b = e[0], e[1]
        label = e[2] if len(e) > 2 else None
        ax_, ay_, aw, ah = nodes[a][0], nodes[a][1], nodes[a][2], nodes[a][3]
        bx_, by_, bw, bh = nodes[b][0], nodes[b][1], nodes[b][2], nodes[b][3]
        p1 = _anchor(ax_, ay_, aw, ah, bx_, by_)
        p2 = _anchor(bx_, by_, bw, bh, ax_, ay_)
        arr = FancyArrowPatch(p1, p2, arrowstyle='-|>', mutation_scale=14,
                              lw=1.6, color=RED, shrinkA=0, shrinkB=0,
                              connectionstyle='arc3,rad=0.0', zorder=1)
        ax.add_patch(arr)
        if label:
            mx, my = (p1[0]+p2[0])/2, (p1[1]+p2[1])/2
            ax.text(mx, my, label, fontsize=7.5, color=INK, ha='center', va='center',
                    bbox=dict(boxstyle='round,pad=0.18', fc='white', ec='#cccccc', lw=0.6), zorder=3)

    # nodes
    for name, (cx, cy, w, h, label, kind) in nodes.items():
        if kind == 'decision':
            pts = [(cx, cy+h/2), (cx+w/2, cy), (cx, cy-h/2), (cx-w/2, cy)]
            poly = Polygon(pts, closed=True, fc='white', ec=RED, lw=1.8, zorder=2)
            ax.add_patch(poly)
            ax.text(cx, cy, label, ha='center', va='center', fontsize=7.6, color=INK, zorder=3)
        elif kind == 'terminal':
            box = FancyBboxPatch((cx-w/2, cy-h/2), w, h, boxstyle='round,pad=0.02,rounding_size=0.18',
                                 fc=RED, ec=RED, lw=1.5, zorder=2)
            ax.add_patch(box)
            ax.text(cx, cy, label, ha='center', va='center', fontsize=8.2, color='white',
                    fontweight='bold', zorder=3)
        elif kind == 'group':
            box = FancyBboxPatch((cx-w/2, cy-h/2), w, h, boxstyle='round,pad=0.02,rounding_size=0.10',
                                 fc='none', ec=RED, lw=1.4, ls='--', zorder=2)
            ax.add_patch(box)
            ax.text(cx, cy+h/2-0.16, label, ha='center', va='top', fontsize=8.4, color=RED,
                    fontweight='bold', zorder=3)
        else:  # process
            box = FancyBboxPatch((cx-w/2, cy-h/2), w, h, boxstyle='round,pad=0.02,rounding_size=0.12',
                                 fc=CHARCOAL, ec=PANEL, lw=1.2, zorder=2)
            ax.add_patch(box)
            ax.text(cx, cy, label, ha='center', va='center', fontsize=8.0, color='white', zorder=3)

    xs = [v[0] for v in nodes.values()]; ys = [v[1] for v in nodes.values()]
    ws = [v[2] for v in nodes.values()]; hs = [v[3] for v in nodes.values()]
    ax.set_xlim(min(x-w/2 for x,w in zip(xs,ws))-0.4, max(x+w/2 for x,w in zip(xs,ws))+0.4)
    ax.set_ylim(min(y-h/2 for y,h in zip(ys,hs))-0.4, max(y+h/2 for y,h in zip(ys,hs))+0.4)
    ax.set_aspect('equal'); ax.axis('off')
    if title:
        ax.set_title(title, fontsize=11, color=INK, fontweight='bold', pad=10)
    fig.savefig(path, bbox_inches='tight', facecolor=PAGE_BG, dpi=200)
    plt.close(fig)
    return path

# =====================================================================
# BUILD DIAGRAMS
# =====================================================================
def P(cx,cy,label,w=2.6,h=0.9,kind='process'): return (cx,cy,w,h,label,kind)

# --- 1. Master end-to-end journey (horizontal snake, 2 rows x 5) ---
n = {}
labels = ['Visitor /\nLead','Register','Verify\nEmail','Login','KYC',
          'Open\nAccount','Deposit','Trade','Withdraw','Support\n& IB']
BW, BH, GX, ROWY = 2.3, 1.2, 3.0, -2.4
order=[]
for i,l in enumerate(labels):
    if i < 5:                      # top row, left -> right
        cx, cy = i*GX, 0
    else:                          # bottom row, right -> left
        cx, cy = (9-i)*GX, ROWY
    kind = 'terminal' if i in (0,9) else 'process'
    nm=f'j{i}'; n[nm]=P(cx, cy, l, w=BW, h=BH, kind=kind); order.append(nm)
edges=[(order[i],order[i+1]) for i in range(len(order)-1)]
draw_flow(n, edges, os.path.join(DIAG,'d1_journey.png'), (13.0, 4.0))

# --- 2. Lead-to-client conversion (vertical with statuses) ---
n={}
seq=[('Lead Captured','status: New'),('Nurtured / Contacted','status: Contacted'),
     ('Registers Live Account','status: Pending'),('Email + KYC Verified','status: Verified'),
     ('First Deposit Cleared','status: Funded'),('Active Client','status: Active')]
y=0; order=[]
for i,(l,s) in enumerate(seq):
    nm=f'c{i}'; kind='terminal' if i in (0,len(seq)-1) else 'process'
    n[nm]=P(0,y,l,w=3.0,kind=kind); order.append((nm,s)); y-=1.4
edges=[(order[i][0],order[i+1][0], order[i+1][1]) for i in range(len(order)-1)]
draw_flow(n, edges, os.path.join(DIAG,'d2_conversion.png'), (5.2, 8.6))

# --- 3. Registration & onboarding (with verify decision) ---
n={
 'a':P(0,0,'Start Registration',kind='terminal'),
 'b':P(0,-1.4,'Step 1: Personal Details'),
 'c':P(0,-2.8,'Step 2: Account Preferences'),
 'd':P(0,-4.2,'Step 3: Credentials + T&C'),
 'e':P(0,-5.6,'Create User (Pending)'),
 'f':(0,-7.1,2.6,1.3,'Email\nVerified?','decision'),
 'g':P(3.2,-7.1,'Resend Email',w=2.2),
 'h':P(0,-8.6,'Login + Onboarding',kind='terminal'),
}
edges=[('a','b'),('b','c'),('c','d'),('d','e'),('e','f'),
       ('f','g','No'),('f','h','Yes')]
draw_flow(n, edges, os.path.join(DIAG,'d3_registration.png'), (6.2, 9.8))

# --- 4. KYC verification (steps + review decision + reject loop) ---
n={
 'a':P(0,0,'Start KYC',kind='terminal'),
 'id':P(0,-1.4,'Upload Identity Doc'),
 'ad':P(0,-2.8,'Upload Proof of Address'),
 'sf':P(0,-4.2,'Upload Selfie + ID'),
 'pend':P(0,-5.6,'Status: Pending Review'),
 'rev':(0,-7.2,2.8,1.4,'Compliance\nApproves?','decision'),
 'rej':P(3.4,-7.2,'Rejected:\nResubmit',w=2.4),
 'ok':P(0,-8.8,'KYC Verified',kind='terminal'),
 'unlock':P(0,-10.1,'Withdrawals Unlocked'),
}
edges=[('a','id'),('id','ad'),('ad','sf'),('sf','pend'),('pend','rev'),
       ('rev','rej','No'),('rev','ok','Yes'),('ok','unlock')]
draw_flow(n, edges, os.path.join(DIAG,'d4_kyc.png'), (6.4, 11.4))

# --- 5. Deposit flow ---
n={
 'a':P(0,0,'Select Deposit Method',kind='terminal'),
 'b':P(0,-1.4,'Enter Amount + Account'),
 'c':P(0,-2.8,'Confirm (Pending)'),
 'd':(0,-4.3,2.8,1.4,'Payment\nSuccess?','decision'),
 'e':P(3.4,-4.3,'Failed:\nRetry',w=2.2),
 'f':P(0,-5.8,'Processing'),
 'g':P(0,-7.2,'Completed: Balance Credited',w=3.2,kind='terminal'),
}
edges=[('a','b'),('b','c'),('c','d'),('d','e','No'),('d','f','Yes'),('f','g')]
draw_flow(n, edges, os.path.join(DIAG,'d5_deposit.png'), (6.4, 8.6))

# --- 6. Withdrawal flow ---
n={
 'a':P(0,0,'Request Withdrawal',kind='terminal'),
 'b':(0,-1.6,3.0,1.4,'KYC + Funds\nValid?','decision'),
 'b2':P(3.7,-1.6,'Blocked:\nComplete KYC',w=2.4),
 'c':P(0,-3.2,'Pending (Funds Held)'),
 'd':(0,-4.8,3.0,1.4,'Finance\nApproves?','decision'),
 'e':P(3.7,-4.8,'Rejected:\nHold Released',w=2.6),
 'f':P(0,-6.4,'Paid / Completed',kind='terminal'),
}
edges=[('a','b'),('b','b2','No'),('b','c','Yes'),('c','d'),
       ('d','e','No'),('d','f','Yes')]
draw_flow(n, edges, os.path.join(DIAG,'d6_withdrawal.png'), (6.6, 8.2))

# --- 7. System actors & modules map ---
n={
 'pub':(-3.4, 1.6, 3.0, 1.2, 'PUBLIC WEBSITE\nHome · Markets · Accounts\nPartnership · Contact','group'),
 'auth':(0, 1.6, 2.6, 1.2, 'ONBOARDING\nRegister · Demo\nEmail Verify','group'),
 'portal':(3.4, 1.6, 3.0, 1.2, 'CLIENT PORTAL\nDashboard · Accounts\nFunds · KYC · Support','group'),
 'admin':(-3.4, -1.4, 3.0, 1.2, 'ADMIN BACK-OFFICE\nUsers · Compliance\nFinance · CMS','group'),
 'ib':(0, -1.4, 2.6, 1.2, 'PARTNER / IB\nReferrals · Rebates\nReports','group'),
 'sys':(3.4, -1.4, 3.0, 1.2, 'SYSTEM\nNotifications · Audit\nRisk · Scheduler','group'),
}
edges=[('pub','auth'),('auth','portal'),('portal','ib'),('admin','portal'),('sys','portal')]
draw_flow(n, edges, os.path.join(DIAG,'d7_modules.png'), (8.2, 5.2))

# =====================================================================
# HORIZONTAL (compact) variants — for inline use in the workflow doc
# =====================================================================
GX, MW, MH, BRY = 2.95, 2.5, 1.15, -2.35
def hnode(i, label, kind='process', row=0):
    return P(i*GX, 0 if row==0 else BRY, label, w=MW, h=MH, kind=kind)
def hdiamond(i, label, row=0):
    return (i*GX, 0 if row==0 else BRY, 2.7, 1.35, label, 'decision')

# h2 conversion (linear)
labels=['Lead\nCaptured','Nurtured','Registered\n(Pending)','Verified','Funded','Active\nClient']
n={f'c{i}':hnode(i,l,'terminal' if i in (0,5) else 'process') for i,l in enumerate(labels)}
edges=[(f'c{i}',f'c{i+1}') for i in range(5)]
draw_flow(n, edges, os.path.join(DIAG,'h2_conversion.png'), (15.5, 2.4))

# h3 registration (6 main + 1 branch)
n={'a':hnode(0,'Personal\nDetails'),'b':hnode(1,'Account\nPreferences'),
   'c':hnode(2,'Credentials\n+ T&C'),'d':hnode(3,'Create User\n(Pending)'),
   'e':hdiamond(4,'Email\nVerified?'),'f':hnode(5,'Login +\nOnboarding','terminal'),
   'g':hnode(4,'Resend\nEmail',row=1)}
edges=[('a','b'),('b','c'),('c','d'),('d','e'),('e','f','Yes'),('e','g','No')]
draw_flow(n, edges, os.path.join(DIAG,'h3_registration.png'), (15.5, 4.3))

# h4 kyc (5 main + 1 branch)
n={'a':hnode(0,'Upload\nIdentity'),'b':hnode(1,'Proof of\nAddress'),
   'c':hnode(2,'Selfie\n+ ID'),'d':hdiamond(3,'Compliance\nApproves?'),
   'e':hnode(4,'KYC Verified\n(Unlocks)','terminal'),
   'r':hnode(3,'Rejected:\nResubmit',row=1)}
edges=[('a','b'),('b','c'),('c','d'),('d','e','Yes'),('d','r','No')]
draw_flow(n, edges, os.path.join(DIAG,'h4_kyc.png'), (13.5, 4.3))

# h5 deposit (5 main + 1 branch)
n={'a':hnode(0,'Select\nMethod'),'b':hnode(1,'Enter\nAmount'),
   'c':hnode(2,'Confirm\n(Pending)'),'d':hdiamond(3,'Payment\nSuccess?'),
   'e':hnode(4,'Completed:\nCredited','terminal'),
   'f':hnode(3,'Failed:\nRetry',row=1)}
edges=[('a','b'),('b','c'),('c','d'),('d','e','Yes'),('d','f','No')]
draw_flow(n, edges, os.path.join(DIAG,'h5_deposit.png'), (13.5, 4.3))

# h6 withdrawal (5 main + 2 branches)
n={'a':hnode(0,'Request','terminal'),'b':hdiamond(1,'KYC + Funds\nValid?'),
   'c':hnode(2,'Pending\n(Held)'),'d':hdiamond(3,'Finance\nApproves?'),
   'e':hnode(4,'Paid /\nCompleted','terminal'),
   'b2':hnode(1,'Blocked:\nComplete KYC',row=1),'e2':hnode(3,'Rejected:\nHold Released',row=1)}
edges=[('a','b'),('b','c','Yes'),('b','b2','No'),('c','d'),('d','e','Yes'),('d','e2','No')]
draw_flow(n, edges, os.path.join(DIAG,'h6_withdrawal.png'), (13.5, 4.3))

print('Diagrams rendered to', DIAG)

# =====================================================================
# BUILD DOCX
# =====================================================================
doc = Document()
nm = doc.styles['Normal']; nm.font.name='Calibri'; nm.font.size=Pt(10.5)
nm.paragraph_format.space_after=Pt(6)
for s in doc.sections:
    s.top_margin=Inches(0.8); s.bottom_margin=Inches(0.8)
    s.left_margin=Inches(0.9); s.right_margin=Inches(0.9)

def shade(cell,hexc):
    tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),hexc); tcPr.append(shd)
def H(text,level=1):
    p=doc.add_heading(text,level=level)
    for r in p.runs: r.font.color.rgb = RED_RGB if level==1 else RGBColor(0x11,0x11,0x11)
    return p
def para(text,bold=False,italic=False,color=None,size=None,after=6):
    p=doc.add_paragraph(); r=p.add_run(text); r.bold=bold; r.italic=italic
    if color:r.font.color.rgb=color
    if size:r.font.size=Pt(size)
    p.paragraph_format.space_after=Pt(after); return p
def step(text): return doc.add_paragraph(text, style='List Number')
def bullet(text,prefix=None):
    p=doc.add_paragraph(style='List Bullet')
    if prefix: rr=p.add_run(prefix); rr.bold=True; p.add_run(text)
    else: p.add_run(text)
    return p
def img(path,width=6.3):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Inches(width)); return p
def caption(text):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(text); r.italic=True; r.font.size=Pt(8.5); r.font.color.rgb=GRAY_RGB; return p
def table(headers,rows,widths=None):
    t=doc.add_table(rows=1,cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style='Light Grid Accent 1'
    for i,head in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=''; run=c.paragraphs[0].add_run(head)
        run.bold=True; run.font.color.rgb=WHITE_RGB; run.font.size=Pt(9); shade(c,'E11D2E')
    for row in rows:
        cells=t.add_row().cells
        for i,v in enumerate(row):
            cells[i].text=''; run=cells[i].paragraphs[0].add_run(str(v)); run.font.size=Pt(9)
    if widths:
        for i,w in enumerate(widths):
            for r_ in t.rows: r_.cells[i].width=Inches(w)
    doc.add_paragraph(); return t
def set_valign(cell, val='center'):
    tcPr=cell._tc.get_or_add_tcPr(); v=OxmlElement('w:vAlign'); v.set(qn('w:val'),val); tcPr.append(v)
def no_split(row):
    trPr=row._tr.get_or_add_trPr(); cs=OxmlElement('w:cantSplit'); trPr.append(cs)
def side_by_side(image_path, img_w, lines, img_left=True):
    """Two-column borderless block: diagram on one side, bullet narrative on the other."""
    t=doc.add_table(rows=1, cols=2); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False
    no_split(t.rows[0])
    ci, ct = (0,1) if img_left else (1,0)
    cimg=t.rows[0].cells[ci]; ctxt=t.rows[0].cells[ct]
    cimg.width=Inches(3.7); ctxt.width=Inches(2.8)
    set_valign(cimg); set_valign(ctxt)
    p=cimg.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(image_path, width=Inches(img_w))
    first=True
    for ln in lines:
        p = ctxt.paragraphs[0] if first else ctxt.add_paragraph(); first=False
        rb=p.add_run('▸ '); rb.font.color.rgb=RED_RGB; rb.bold=True
        p.add_run(ln).font.size=Pt(9.5)
        p.paragraph_format.space_after=Pt(5)
    doc.add_paragraph(); return t
def field(p,code):
    for t,attr in [('begin',None),('inst',code),('separate',None),('text','?'),('end',None)]:
        run=p.add_run()
        if t=='inst':
            el=OxmlElement('w:instrText'); el.set(qn('xml:space'),'preserve'); el.text=attr; run._r.append(el)
        elif t=='text': run.text=attr
        else:
            el=OxmlElement('w:fldChar'); el.set(qn('w:fldCharType'),t); run._r.append(el)

# footer page numbers
fp=doc.sections[0].footer.paragraphs[0]; fp.alignment=WD_ALIGN_PARAGRAPH.CENTER
fp.add_run('27 Markets — How the Platform Works   |   Page ').font.size=Pt(8)
field(fp,'PAGE'); fp.add_run(' of ').font.size=Pt(8); field(fp,'NUMPAGES')

# ---- COVER ----
t=doc.add_paragraph(); t.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=t.add_run('27 MARKETS'); r.bold=True; r.font.size=Pt(36); r.font.color.rgb=RED_RGB
s=doc.add_paragraph(); s.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=s.add_run('Trade Beyond Limits'); r.italic=True; r.font.size=Pt(13); r.font.color.rgb=GRAY_RGB
doc.add_paragraph()
s2=doc.add_paragraph(); s2.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=s2.add_run('How the Platform Works'); r.bold=True; r.font.size=Pt(20)
s3=doc.add_paragraph(); s3.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=s3.add_run('End-to-End Workflow — Client Overview'); r.font.size=Pt(13); r.font.color.rgb=GRAY_RGB
doc.add_paragraph()
m=doc.add_paragraph(); m.alignment=WD_ALIGN_PARAGRAPH.CENTER
m.add_run('Prepared for: Client   |   Date: 2026-06-22   |   Version 1.0').font.size=Pt(10)
doc.add_paragraph()
i=doc.add_paragraph(); i.alignment=WD_ALIGN_PARAGRAPH.CENTER
ir=i.add_run('This document explains, in plain language and visual diagrams, how the 27 Markets trading '
             'platform works from a visitor’s first click through to active, funded trading — including '
             'onboarding, identity verification, funding, withdrawals, and ongoing support.')
ir.italic=True; ir.font.color.rgb=GRAY_RGB; ir.font.size=Pt(10.5)
doc.add_page_break()

# ---- TOC ----
th=doc.add_paragraph(); r=th.add_run('Contents'); r.bold=True; r.font.size=Pt(16); r.font.color.rgb=RED_RGB
tp=doc.add_paragraph(); field(tp,'TOC \\o "1-1" \\h \\z \\u')
doc.add_page_break()

# ---- 1. OVERVIEW ----
H('1. Platform at a Glance',1)
para('27 Markets is a premium online brokerage made of two connected parts that work as one product:')
bullet('A public marketing website that introduces the brand, markets, account types, and partnership program, and converts visitors into registered clients.', 'Public Website — ')
bullet('A secure client portal where verified clients open accounts, deposit and withdraw funds, complete identity checks, download platforms, and get support.', 'Client Portal — ')
para('Behind the scenes, an admin back-office and automated system handle compliance reviews, finance approvals, notifications, and partner rebates.')
img(os.path.join(DIAG,'d7_modules.png'), width=6.5)
caption('Figure 1 — The platform modules and how they connect.')

# ---- 2. MASTER JOURNEY ----
H('2. End-to-End Client Journey',1)
para('The diagram below shows the complete path a person takes — from first discovering 27 Markets to '
     'actively trading and growing with us. Read it left-to-right across the top row, then right-to-left along '
     'the bottom row.')
img(os.path.join(DIAG,'d1_journey.png'), width=6.6)
caption('Figure 2 — The end-to-end journey, stage by stage.')

# ---- 3. CONVERSION ----
H('3. Becoming a Client (Lead → Active)',1)
para('Every visitor is a potential client. The platform tracks each person’s progress through clear stages, '
     'so the team always knows where someone is in their journey.')
side_by_side(os.path.join(DIAG,'d2_conversion.png'), 3.3, [
    'A lead is captured from the website, a demo request, or a partner referral.',
    'Nurtured with welcome and follow-up communication.',
    'On registration the person becomes a Pending client.',
    'Email + KYC verification moves them to Verified.',
    'The first cleared deposit makes them a Funded, Active client.',
])
caption('Figure 3 — Lead-to-client conversion with status at each stage.')

# ---- 4. REGISTRATION ----
H('4. Registration & Onboarding',1)
para('Signing up is a guided, multi-step form. Each step is validated before the user can continue, and the '
     'account is activated only after the email is verified.')
side_by_side(os.path.join(DIAG,'d3_registration.png'), 3.2, [
    'Step 1 collects personal details (name, email, country).',
    'Step 2 captures account preferences (type, currency, leverage).',
    'Step 3 sets credentials and accepts the terms.',
    'The account is created in a Pending state.',
    'A verification email must be confirmed before first login.',
    'Once verified, the client lands in guided onboarding.',
], img_left=False)
caption('Figure 4 — Registration and onboarding flow.')

# ---- 5. KYC ----
H('5. Identity Verification (KYC)',1)
para('To keep the platform secure and compliant, clients verify their identity in three steps. A compliance '
     'officer reviews each submission. Withdrawals are unlocked once verification is approved.')
side_by_side(os.path.join(DIAG,'d4_kyc.png'), 3.2, [
    'Upload an identity document (ID or passport).',
    'Upload proof of address (utility bill or statement).',
    'Upload a selfie holding the ID.',
    'Submissions move to Pending Review.',
    'Compliance approves, or rejects with a reason to resubmit.',
    'Once approved, KYC is Verified and withdrawals unlock.',
])
caption('Figure 5 — KYC verification with compliance review.')

# ---- 6. DEPOSIT ----
H('6. Depositing Funds',1)
para('Clients can fund their account using several methods — bank transfer, card, e-wallet, or crypto. '
     'Most methods are instant; bank transfers take 1–3 business days.')
side_by_side(os.path.join(DIAG,'d5_deposit.png'), 3.3, [
    'Choose a deposit method.',
    'Enter the amount and the account to credit.',
    'Confirm — the transaction is created as Pending.',
    'If the payment succeeds it moves to Processing.',
    'On completion the account balance is credited instantly.',
    'A failed payment keeps the balance unchanged and can be retried.',
], img_left=False)
caption('Figure 6 — Deposit flow from method selection to credited balance.')

# ---- 7. WITHDRAWAL ----
H('7. Withdrawing Funds',1)
para('Withdrawals are protected by identity and balance checks, then approved by the finance team. Funds are '
     'returned to the original payment source for security.')
side_by_side(os.path.join(DIAG,'d6_withdrawal.png'), 3.4, [
    'The client requests a withdrawal.',
    'The system checks KYC status and available balance.',
    'If valid, funds are held and the request goes Pending.',
    'The finance team reviews and approves the payout.',
    'Approved funds are Paid to the original source.',
    'A rejection releases the hold and restores the balance.',
])
caption('Figure 7 — Withdrawal flow with verification and finance approval.')

# ---- 8. STATUS REFERENCE ----
H('8. Status Reference',1)
para('Every important item in the platform has a clear status, so clients and staff always know what is '
     'happening.')
table(['Area','Possible statuses'], [
    ['Client account','Pending, Verified, Active, Suspended, Closed'],
    ['KYC step','Not Submitted, Pending, Approved, Rejected'],
    ['Trading account','Active, Inactive, Suspended'],
    ['Deposit','Pending, Processing, Completed, Failed'],
    ['Withdrawal','Pending, Approved, Paid, Rejected'],
    ['Support ticket','Open, In Progress, Resolved, Closed'],
], widths=[1.9,4.6])

# ---- 9. IMPORTANT NOTE: COSTS & GO-LIVE ----
H('9. Important Note — Costs & Go-Live',1)
para('Please note the following before launch:', bold=True)
bullet('This version uses simulated data for demonstration. Live trading is not active yet.', 'Current stage: ')
bullet('Real trading starts only after the trading platform, payment gateways, identity verification, market-data feed, and brokerage licensing are connected and approved.', 'Going live: ')
bullet('These are paid third-party services. Their API keys carry setup and ongoing usage charges (billed by each provider) and are not free of cost or included in this build.', 'Third-party costs: ')
bullet('Final pricing depends on the chosen providers, volume, and region, and will be shared in a separate quotation.', 'Pricing: ')

# ---- 10. SUMMARY ----
H('10. Summary',1)
para('27 Markets delivers a complete, secure brokerage experience: a polished public website that earns trust '
     'and converts visitors, and a full client portal for verification, funding, trading, and support — all '
     'governed by clear statuses, compliance checks, and finance controls. The result is a smooth journey for '
     'clients and full operational visibility for the business.')

if os.path.exists(LOGO):
    make_watermark(LOGO, WM)
    set_watermark(doc.sections[0], WM)

out = os.path.join(BASE, '27-Markets-How-It-Works.docx')
try:
    doc.save(out)
    print('Saved:', out)
except PermissionError:
    alt = os.path.join(BASE, '27-Markets-How-It-Works-v2.docx')
    doc.save(alt)
    print('Original file is open/locked. Saved to:', alt)
